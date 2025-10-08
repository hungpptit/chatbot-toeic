# Script đọc file DOCX và extract text, images, tables
import os
import zipfile
import shutil
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
except ImportError:
    print("❌ Cần cài đặt python-docx:")
    print("   pip install python-docx")
    print("   pip install pillow")
    exit(1)

def extract_images_from_docx(docx_path, output_folder='extracted_images'):
    """Extract tất cả images từ file DOCX"""
    
    # Tạo thư mục output
    output_path = Path(output_folder)
    output_path.mkdir(exist_ok=True)
    
    # DOCX file thực chất là ZIP archive
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        # Tìm tất cả file image trong word/media/
        image_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
        
        print(f"\n📸 Tìm thấy {len(image_files)} hình ảnh:")
        
        for i, img_file in enumerate(image_files, 1):
            # Extract image
            img_name = os.path.basename(img_file)
            output_file = output_path / f"image_{i}_{img_name}"
            
            with zip_ref.open(img_file) as source, open(output_file, 'wb') as target:
                shutil.copyfileobj(source, target)
            
            print(f"  ✅ {output_file}")
    
    return len(image_files)

def read_docx_content(docx_path):
    """Đọc nội dung text, tables từ DOCX"""
    
    doc = Document(docx_path)
    
    content = []
    content.append("=" * 80)
    content.append("📄 NỘI DUNG FILE: " + os.path.basename(docx_path))
    content.append("=" * 80 + "\n")
    
    # Đọc từng element (paragraph hoặc table)
    for element in doc.element.body:
        if isinstance(element, CT_P):
            # Paragraph
            para = Paragraph(element, doc)
            text = para.text.strip()
            
            if text:
                # Check style (với error handling)
                try:
                    style = para.style.name if para.style else "Normal"
                except:
                    style = "Normal"
                
                if 'Heading' in style:
                    # Heading
                    level = style.replace('Heading ', '').strip()
                    content.append("\n" + "#" * min(int(level) if level.isdigit() else 1, 6) + " " + text)
                else:
                    # Normal paragraph
                    content.append(text)
        
        elif isinstance(element, CT_Tbl):
            # Table
            table = Table(element, doc)
            content.append("\n📊 TABLE:")
            content.append("-" * 80)
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                content.append(" | ".join(row_data))
            
            content.append("-" * 80)
    
    return "\n".join(content)

def analyze_docx_structure(docx_path):
    """Phân tích cấu trúc document"""
    
    doc = Document(docx_path)
    
    stats = {
        'paragraphs': len(doc.paragraphs),
        'tables': len(doc.tables),
        'sections': len(doc.sections),
        'styles_used': set()
    }
    
    # Thu thập styles được sử dụng
    for para in doc.paragraphs:
        if para.style:
            stats['styles_used'].add(para.style.name)
    
    print("\n📊 THỐNG KÊ DOCUMENT:")
    print(f"  - Số đoạn văn: {stats['paragraphs']}")
    print(f"  - Số bảng: {stats['tables']}")
    print(f"  - Số sections: {stats['sections']}")
    print(f"  - Styles sử dụng: {', '.join(sorted(stats['styles_used']))}")
    
    return stats

def main():
    docx_file = 'd:/Chatbot_Toeic/BaoCaoHT.docx'
    
    if not os.path.exists(docx_file):
        print(f"❌ Không tìm thấy file: {docx_file}")
        return
    
    print(f"🔍 Đang đọc file: {docx_file}\n")
    
    # 1. Extract images
    try:
        img_count = extract_images_from_docx(docx_file, 'd:/Chatbot_Toeic/BaoCaoHT_images')
    except Exception as e:
        print(f"⚠️ Lỗi khi extract images: {e}")
        img_count = 0
    
    # 2. Analyze structure
    try:
        stats = analyze_docx_structure(docx_file)
    except Exception as e:
        print(f"⚠️ Lỗi khi phân tích structure: {e}")
    
    # 3. Read content
    try:
        content = read_docx_content(docx_file)
        
        # Save to text file
        output_file = 'd:/Chatbot_Toeic/BaoCaoHT_extracted.txt'
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(f"\n✅ Đã lưu nội dung text vào: {output_file}")
        
        # Print preview
        print("\n" + "=" * 80)
        print("📖 PREVIEW (100 dòng đầu):")
        print("=" * 80)
        lines = content.split('\n')
        for line in lines[:100]:
            print(line)
        
        if len(lines) > 100:
            print(f"\n... (còn {len(lines) - 100} dòng nữa)")
    
    except Exception as e:
        print(f"❌ Lỗi khi đọc content: {e}")
    
    print("\n" + "=" * 80)
    print("🎉 HOÀN THÀNH!")
    print("=" * 80)
    print(f"📁 Images: d:/Chatbot_Toeic/BaoCaoHT_images/")
    print(f"📄 Text content: d:/Chatbot_Toeic/BaoCaoHT_extracted.txt")

if __name__ == "__main__":
    main()
